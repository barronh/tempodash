__all__ = ['from_antype']
import pandas as pd
import time
from .cfg import v1start, v2start, v3start
# from docx.enum.section import WD_ORIENT
# from docx.enum.style import WD_STYLE

versionstrs = []
if not pd.isnull(v1start):
    versionstrs.append(
        f' Version 1 (v1) data is from {v1start:%F} to {v2start:%F}.'
    )
if not pd.isnull(v2start):
    versionstrs.append(
        f' Version 2 (v2) includes data both before {v1start:%F} and '
        + f' after {v2start:%F}.'
    )
if not pd.isnull(v3start):
    versionstrs.append(
        f' Version 3 (v3) includes data after {v3start:%F}.'
    )

versionstr = ''.join(versionstrs)

intro = [
    (
        'This report has plots organized by analysis region for {antype}'
        + ' {spc}. TEMPO data over the analysis time-period varies in the'
        + ' algorithm used for the retrieval. ' + versionstr
        + ' All TEMPO data has been filtered for effective cloud fraction less'
        + ' than 0.15 based on SAO recommendation at the GeoXO ACX May'
        + ' meeting.',
        None, None
    ),
    (
        'TropOMI intersections are any TropOMI pixel that overlaps a TEMPO'
        + ' pixel. Statistics (mean, std, quantiles) are all based on'
        + ' all intersections within the region.', None, 'tropomi'
    ),
    (
        'Pandora intersections are any TEMPO pixel that overlaps a 0.03 degree'
        + ' radius. Only intersections within 15 minutes of a TEMPO'
        + ' measurement are used. Statistics (mean, std, quantiles) are all'
        + ' based on each +-15 minute period or pooled data.',
        None, 'pandora'
    ),
    (
        'AirNow intersections are any TEMPO pixel that overlaps a monitor'
        + ' location. AirNow measurements integrated over an hour, and are'
        + ' compared to any TEMPO measurement within that hour.'
        + ' Statistics (mean, std, quantiles) are all based on hourly pairs'
        + ' or pairs pooled within the Ozone non-attainment bounding box.',
        None, 'airnow'
    ),
    (
        '{antype} {spc} Spatial Overview includes scatter plots where each'
        + ' point represents an analysis region and the error bars'
        + ' represent the interquartile range of TEMPO and {antype}.'
        + ' Scatter plots are shown for the whole record (v1+v2), v1, v2,'
        + ' and v3',
        'List Bullet', None
    ),
    (
        '{antype} {spc} Spatial Overview includes maps of the medians (TEMPO'
        + ' and {antype}) and normalized median bias for each analysis region.'
        + ' Maps show the whole data record (v1+v2), v1, v2, and v3'
        + ' separately.',
        'List Bullet', None
    ),
    (
        'Analysis Region Overview includes boxplots of {spc} observations from'
        + ' TEMPO (red) and {antype} (black) by analysis region. The results'
        + ' are showns separately for the whole record (v1+v2), v1, v2, and'
        + ' v3. These are followed by normalized error plots for the same'
        + ' regions and versions.',
        'List Bullet', None
    ),
    (
        'For TropOMI, the Analysis Region Overview has twp sections'
        + ' The first shows TropOMI {spc} at Pandora sites (within +-0.1'
        + ' degree). The second shows TropOMI {spc} at all other analysis'
        + ' regions.', 'List Bullet', 'tropomi'
    ),
    (
        '{antype} {spc} Detail by Analysis Region includes detail plots for'
        + ' each region.',
        'List Bullet', None
    ),
    (
        'Scatter plots are shown for the whole record (v1+v2), v1, v2, and v3'
        + ' include linear regression analyses. (deming needs to be added)',
        'List Bullet 2', None
    ),
    (
        'Observations diurnal patterns are shown for v1+v2, v1, v2, and v3'
        + ' shown by site', 'List Bullet 2', 'pandora'
    ),
    (
        'Observations diurnal patterns are shown for v1+v2, v1, v2, and v3'
        + ' shown by site', 'List Bullet 2', 'airnow'
    ),
    (
        'Hourly time series of all observations are shown by site',
        'List Bullet 2', None
    ),
    (
        'For AirNow, scatter plots used data with normalized by removing'
        + ' the mean and dividing by the standard deviation',
        'List Bullet 2', 'airnow'
    ),
    (
        'Use Headings to navigate directly to a specific analysis region.',
        None, None
    ),
]


def new_doc(antype, spc):
    from docx import Document
    from docx.shared import Inches

    datestr = time.strftime('%Y-%m-%dT%H%z')
    doc = Document()
    sec = doc.sections[-1]
    # new_width, new_height = sec.page_height, sec.page_width
    # sec.orientation = WD_ORIENT.LANDSCAPE
    # sec.page_width = new_width
    # sec.page_height = new_height
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    antype = {
        'airnow': 'AirNow', 'tropomi': 'TropOMI', 'pandora': 'Pandora'
    }.get(antype, antype)
    title = f'{antype} vs TEMPO {spc.upper()}'
    author = 'Barron H. Henderson (US EPA)'
    doc.add_heading(title, 0)
    ap = doc.add_paragraph
    ap(f'generated by: {author}', 'List Paragraph')
    ap(f'last updated: {datestr}', 'List Paragraph')
    doc.add_heading('Report Description', 1)
    for txt, sty, typchk in intro:
        if typchk is None or antype.startswith(typchk):
            doc.add_paragraph(
                txt.format(antype=antype, spc=spc.upper()), style=sty
            )

    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.keywords = f'TEMPO; {antype}'
    return doc


def figsort(path):
    suffix = path.split('_')[-1]
    sord = {'scat': 0, 'ts': 1, 'ss': 3, 'ds': 2}.get(suffix[:-4], 3)
    return path[:-len(suffix)], sord


def from_antype(antype, spc):
    import os
    from docx.shared import Inches
    from .cfg import configs as cfgs, queries
    # _versions = ['all', 'v1', 'v2', 'v3']
    _versions = [qkey for qkey, qstr, qlabel in queries]
    if len(_versions) > 1:
        allkey = 'all'
    else:
        allkey = _versions[0]
    pgwidth = Inches(7.5)
    sumwidth = pgwidth
    mapwidth = pgwidth
    scatwidth = pgwidth / len(_versions) - Inches(0.025)
    doc = new_doc(antype, spc)
    sec = doc.sections[-1]
    w = (sec.page_width - sec.left_margin - sec.right_margin) * 0.98
    # doc.save(f'{antype}.docx')
    # print(bhh)
    doc.add_page_break()
    doc.add_heading(f'{antype} Spatial Overview', 1)
    paragraph = doc.add_paragraph()
    prefix = f'figs/summary/{antype}_{spc}'
    for qkey in _versions:
        run = paragraph.add_run()
        scatpath = f'{prefix}_{qkey}_all_scat.png'
        run.add_picture(scatpath, width=scatwidth)
    if antype != 'airnow':
        for qkey in _versions:
            mappath = f'{prefix}_{qkey}_map.png'
            doc.add_picture(mappath, width=mapwidth)

    if antype.startswith('tropomi'):
        doc.add_heading('Pandora Sites Overview', 1)
        paragraph = doc.add_paragraph()
        for qkey in _versions:
            run = paragraph.add_run()
            sumpath = f'{prefix}_{qkey}_summary_pandora.png'
            run.add_picture(sumpath, width=sumwidth)
        paragraph = doc.add_paragraph()
        for qkey in _versions:
            run = paragraph.add_run()
            sumpath = f'{prefix}_{qkey}_bias_summary_pandora.png'
            run.add_picture(sumpath, width=sumwidth)
        doc.add_heading('Nonattainment Area Regions Overview', 1)
        paragraph = doc.add_paragraph()
        for qkey in _versions:
            run = paragraph.add_run()
            sumpath = f'{prefix}_{qkey}_summary_ozone.png'
            run.add_picture(sumpath, width=sumwidth)
        paragraph = doc.add_paragraph()
        for qkey in _versions:
            run = paragraph.add_run()
            sumpath = f'{prefix}_{qkey}_bias_summary_ozone.png'
            run.add_picture(sumpath, width=sumwidth)
    else:
        if antype == 'pandora':
            doc.add_heading(f'{antype} Overview at Pandora Sites', 1)
        else:
            doc.add_heading(f'{antype} Overview at Analysis Regions', 1)

        paragraph = doc.add_paragraph()
        for qkey in _versions:
            run = paragraph.add_run()
            sumpath = f'{prefix}_{qkey}_summary.png'
            run.add_picture(sumpath, width=sumwidth)
        paragraph = doc.add_paragraph()
        for qkey in _versions:
            run = paragraph.add_run()
            sumpath = f'{prefix}_{qkey}_bias_summary.png'
            run.add_picture(sumpath, width=sumwidth)

    tspath = f'{prefix}_{allkey}_ts.png'
    if os.path.exists(tspath):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(tspath, width=sumwidth)
    dspath = f'{prefix}_{allkey}_all_ds.png'
    if os.path.exists(dspath):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(dspath, width=sumwidth)
    todpath = f'{prefix}_{allkey}_tod_scat.png'
    if os.path.exists(todpath):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(todpath, width=sumwidth)
    monpath = f'{prefix}_{allkey}_month_scat.png'
    if os.path.exists(monpath):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(monpath, width=sumwidth)
    doc.save(f'docs/{antype}_{spc}_summary.docx')
    doc.add_page_break()
    anchktype = antype.replace('_offl', '').replace('_nrti', '')
    doc.add_heading(f'{antype} Detail by Analysis Region', 1)
    first = True
    lockeys = sorted(cfgs, key=lambda x: cfgs[x]['label'])
    for lockey in lockeys:
        cfg = cfgs[lockey]
        loclabel = cfg.get('label', lockey)
        prefix = f'figs/{lockey}/{antype}_{spc}'
        if cfg.get(anchktype, False):
            if first:
                first = False
            else:
                doc.add_page_break()
            doc.add_heading(loclabel, 2)
            paragraph = doc.add_paragraph()
            for qkey in _versions:
                scatpath = f'{prefix}_{qkey}_{lockey}_scat.png'
                if os.path.exists(scatpath):
                    run = paragraph.add_run()
                    run.add_picture(scatpath, width=scatwidth)

            if anchktype not in ('tropomi', 'tropomi_offl', 'tropomi_nrti'):
                for qkey in _versions:
                    dspath = f'{prefix}_{qkey}_{lockey}_ds.png'
                    if os.path.exists(dspath):
                        doc.add_picture(dspath, height=Inches(2.15))

            tspath = f'figs/{lockey}/{antype}_{spc}_{allkey}_{lockey}_ts.png'
            if os.path.exists(tspath):
                doc.add_picture(tspath, width=w)

    doc.save(f'docs/{antype}_{spc}.docx')


if __name__ == '__main__':
    import argparse

    prsr = argparse.ArgumentParser()
    prsr.add_argument('antype', nargs='*')
    args = prsr.parse_args()

    if len(args.antype) == 0:
        args.antype = ['pandora', 'tropomi_offl', 'airnow']

    for antype in args.antype:
        print(f'{antype}...')
        from_antype(antype, 'no2')
        if antype != 'airnow':
            from_antype(antype, 'hcho')
